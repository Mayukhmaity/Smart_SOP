import PyPDF2
import docx
import pandas as pd
import os
import datetime
import stat
import configparser
import re
import fitz

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from pdfminer.layout import LTTextBoxHorizontal
from io import StringIO
from nltk import sent_tokenize


config = configparser.ConfigParser()
config.read('config.ini')

new_file = config.get('indexing', 'filename')

def txt2paragraph(lines):
    # with open(filepath, 'rb') as f:
    #     lines = f.readlines()

    paragraph = ''
    for line in lines:
        if line.isspace():  # is it an empty line?
            if paragraph:
                yield paragraph
                paragraph = ''
            else:
                continue
        else:
            paragraph += ' ' + line.strip()
    yield paragraph
    #f.close()

def paragraph(lines, separator=None):
    if not callable(separator):
        def separator(line): return line == '\n'
    paragraph = []
    for line in lines:
        if separator(line):
            if paragraph:
                yield ''.join(paragraph)
                paragraph = []
        else:
            paragraph.append(line)
    yield ''.join(paragraph)

class WordPdfExtraction:

    def convert_pdf_to_txt(self,path):
        output=[]
        rsrcmgr = PDFResourceManager()
        retstr = StringIO()
        codec = 'utf-8'
        laparams = LAParams()
        device = TextConverter(rsrcmgr, retstr, laparams=laparams)
        fp = open(path, 'rb')
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        password = ""
        maxpages = 0
        caching = True
        pagenos = set()

        # for page in PDFPage.get_pages(fp):
        #     interpreter.process_page(page)
        # # receive the LTPage object for the page.
        # layout = device.get_result()
        # for element in layout:
        #     if instanceof(element, LTTextBoxHorizontal):
        #         print(element.get_text())

        for page_num,page in enumerate(PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password, caching=caching,
                                      check_extractable=True)):

            interpreter.process_page(page)
            text = retstr.getvalue()
            output.append((page_num,text))
            text=''
            retstr.truncate(0)
            retstr.seek(0)

        fp.close()
        device.close()
        retstr.close()
        return output

    #Extract paragraphs from word and convert it into panda dataframe
    def extractWord(self,files):
        this_loc = 1
        df = pd.DataFrame(columns=("name", "paragraph", "content","date_modified"))
        mod_time_dict={}

        for file in files:
            fileStatsObj = os.stat(file)
            modificationTime = datetime.datetime.fromtimestamp(fileStatsObj[stat.ST_MTIME])
            mod_time_dict[file]=modificationTime

        current_time=datetime.datetime.now()

        for file in files:
            last_run_time=datetime.datetime.strptime(config.get('word_indexing','indexing_start_time'),'%Y-%m-%d %H:%M:%S')
            if mod_time_dict[file] < current_time  and mod_time_dict[file] > last_run_time:
                gkzDoc = docx.Document(file)
                paragraphs = gkzDoc.paragraphs
                para_num = 1
                for para in paragraphs:
                    df.loc[this_loc] = file, para_num, para.text,mod_time_dict[file]
                    para_num += 1
                    this_loc += 1
        config.set('word_indexing', 'indexing_start_time',datetime.datetime.strftime(current_time, '%Y-%m-%d %H:%M:%S'))
        # Updating config.ini file
        os.chdir(os.path.dirname(__file__))
        with open('config.ini', 'w') as configfile:
            config.write(configfile)
        return df

    def extractWordSingle(self,file):
        this_loc = 1
        df = pd.DataFrame(columns=("name", "paragraph", "content"))

        gkzDoc = docx.Document(file)
        paragraphs = gkzDoc.paragraphs
        para_num = 1
        for para in paragraphs:
            df.loc[this_loc] = file, para_num, para.text
            para_num += 1
            this_loc += 1
        return df
    # Extract pages from word and convert it into panda dataframe

    def extractPdf(self,files):
        this_loc = 1
        df = pd.DataFrame(columns=("name", "page","para","content"))
        mod_time_dict={}

        for file in files:
            fileStatsObj = os.stat(file)
            modificationTime = datetime.datetime.fromtimestamp(fileStatsObj[stat.ST_MTIME])
            mod_time_dict[file]=modificationTime

        current_time=datetime.datetime.now()

        for file in files:
            last_run_time=datetime.datetime.strptime(config.get('indexing','indexing_start_time'),'%Y-%m-%d %H:%M:%S')
            if mod_time_dict[file] < current_time  and mod_time_dict[file] > last_run_time:
                output=self.convert_pdf_to_txt(file)
                #print(output)
                pagenum=1
                p_num = 1
                for page_num,content in output:
                    paragraphs = content.split('\n\n')
                    #paragraphs= sent_tokenize(content)
                    #paragraphs = re.split('\s{4,}', content)
                    #paragraphs = content.splitlines(True)
                    for p in paragraphs:
                        df.loc[this_loc] = file, page_num, p_num,p.replace("\n", "")
                        this_loc = this_loc + 1
                        p_num+=1
                    pagenum+=1
            #print(output)
        config.set('indexing', 'indexing_start_time',datetime.datetime.strftime(current_time, '%Y-%m-%d %H:%M:%S'))
        # Updating config.ini file
        os.chdir(os.path.dirname(__file__))
        return df